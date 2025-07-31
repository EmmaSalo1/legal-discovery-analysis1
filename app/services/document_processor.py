from datetime import datetime
import os
import mimetypes
import logging
from typing import Dict, Any, List
from app.services.audio_processor import EnhancedAudioProcessor
from app.services.video_processor import EnhancedVideoProcessor
from app.services.image_processor import EnhancedImageProcessor
from app.config import settings

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    def __init__(self):
        self.audio_processor = EnhancedAudioProcessor()
        self.video_processor = EnhancedVideoProcessor()
        self.image_processor = EnhancedImageProcessor()
        
        logger.info("Enhanced Document Processor initialized with multimedia support")
        
    async def process_document(self, file_path: str, case_id: str) -> Dict[str, Any]:
        """Enhanced document processor supporting comprehensive multimedia analysis"""
        
        # Determine file type with better detection
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"Processing file: {file_path} (type: {mime_type}, ext: {file_extension})")
        
        # Route to appropriate processor with correct precedence
        if self._is_audio_file(file_extension, mime_type):
            return await self.audio_processor.process_audio_file(file_path, case_id)
        
        elif self._is_video_file(file_extension, mime_type):
            return await self.video_processor.process_video_file(file_path, case_id)
        
        elif self._is_document_file(file_extension, mime_type):  # Check documents BEFORE images
            return await self._process_traditional_document(file_path, case_id)
        
        elif self._is_image_file(file_extension, mime_type):
            return await self.image_processor.process_image_file(file_path, case_id)
        
        else:
            # Process as traditional document by default
            logger.info(f"Unknown file type, processing as document: {file_path}")
            return await self._process_traditional_document(file_path, case_id)
    
    def _is_audio_file(self, extension: str, mime_type: str) -> bool:
        """Check if file is audio"""
        return (extension in settings.supported_audio_formats or 
                (mime_type and mime_type.startswith('audio/')))
    
    def _is_video_file(self, extension: str, mime_type: str) -> bool:
        """Check if file is video"""
        return (extension in settings.supported_video_formats or 
                (mime_type and mime_type.startswith('video/')))
    
    def _is_image_file(self, extension: str, mime_type: str) -> bool:
        """Check if file is an image (exclude PDFs)"""
        image_extensions = ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']
        return (extension in image_extensions or 
                (mime_type and mime_type.startswith('image/') and 'pdf' not in mime_type.lower()))
    
    def _is_document_file(self, extension: str, mime_type: str) -> bool:
        """Check if file is a document (including PDFs)"""
        document_extensions = ['.pdf', '.docx', '.doc', '.txt', '.rtf']
        return (extension in document_extensions or 
                (mime_type and ('pdf' in mime_type.lower() or 
                              'document' in mime_type.lower() or 
                              'text' in mime_type.lower())))
    
    async def _process_traditional_document(self, file_path: str, case_id: str) -> Dict[str, Any]:
        """Process traditional document formats with enhanced PDF support"""
        try:
            content = ""
            file_extension = os.path.splitext(file_path)[1].lower()
            
            logger.info(f"Processing {file_extension} document: {file_path}")
            
            if file_extension == '.pdf':
                content = await self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = await self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                content = await self._extract_txt_text(file_path)
            elif file_extension == '.rtf':
                content = await self._extract_rtf_text(file_path)
            else:
                # Try to read as text
                content = await self._extract_generic_text(file_path)
            
            # Extract entities and scan for privilege
            entities = await self._extract_document_entities(content)
            privilege_flags = await self._scan_document_privilege(content)
            
            # Analyze document structure
            document_structure = await self._analyze_document_structure(content, file_extension)
            
            # Generate summary
            summary = await self._generate_document_summary(content, entities, document_structure, file_extension)
            
            logger.info(f"Successfully extracted {len(content)} characters from {file_path}")
            
            result = {
                'id': f"doc_{case_id}_{os.path.basename(file_path)}",
                'case_id': case_id,
                'file_path': file_path,
                'file_type': 'document',
                'content': content,
                'entities': entities,
                'privilege_flags': privilege_flags,
                'document_structure': document_structure,
                'summary': summary,
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'file_size': os.path.getsize(file_path),
                    'format': file_extension,
                    'content_length': len(content),
                    'word_count': len(content.split()) if content else 0,
                    'page_count': content.count('--- Page ') if '--- Page ' in content else 1
                },
                'processing_status': 'completed',
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                'id': f"doc_{case_id}_{os.path.basename(file_path)}",
                'case_id': case_id,
                'file_path': file_path,
                'file_type': 'document',
                'error': str(e),
                'processing_status': 'failed',
                'processed_at': datetime.utcnow().isoformat()
            }
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        content = ""
        
        # Method 1: Try PyMuPDF first (most reliable)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            logger.info(f"PDF has {len(doc)} pages (PyMuPDF)")
            
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    
                    if text.strip():
                        content += f"\n--- Page {page_num + 1} ---\n"
                        content += text + "\n"
                    else:
                        # If no text, try OCR on the page image
                        logger.info(f"No text on page {page_num + 1}, trying OCR...")
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        
                        # Save temporary image and use OCR
                        temp_img_path = os.path.join(settings.temp_directory, f"temp_page_{page_num}.png")
                        with open(temp_img_path, "wb") as img_file:
                            img_file.write(img_data)
                        
                        # Use our image processor for OCR
                        ocr_result = await self.image_processor.process_image_file(temp_img_path, "temp")
                        if ocr_result.get('ocr_results', {}).get('combined_text'):
                            content += f"\n--- Page {page_num + 1} (OCR) ---\n"
                            content += ocr_result['ocr_results']['combined_text'] + "\n"
                        
                        # Clean up temp file
                        if os.path.exists(temp_img_path):
                            os.unlink(temp_img_path)
                            
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1} with PyMuPDF: {e}")
            
            doc.close()
            
            if content.strip():
                logger.info(f"PyMuPDF extracted {len(content)} characters")
                return content.strip()
                
        except ImportError:
            logger.warning("PyMuPDF not available - install with: pip install pymupdf")
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"PDF has {len(pdf.pages)} pages (pdfplumber)")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1} with pdfplumber: {e}")
                        
            if content.strip():
                logger.info(f"pdfplumber extracted {len(content)} characters")
                return content.strip()
                
        except ImportError:
            logger.warning("pdfplumber not available - install with: pip install pdfplumber")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 3: Try PyPDF2 as final fallback
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                logger.info(f"PDF has {len(pdf_reader.pages)} pages (PyPDF2)")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        
            if content.strip():
                logger.info(f"PyPDF2 extracted {len(content)} characters")
                return content.strip()
                
        except ImportError:
            logger.warning("PyPDF2 not available")
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # If all methods fail
        if not content.strip():
            logger.error(f"Failed to extract text from PDF: {file_path}")
            return f"Error: Could not extract text from PDF file {os.path.basename(file_path)}. This may be a scanned PDF requiring OCR processing."
        
        return content.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            # Extract headers and footers if possible
            try:
                for section in doc.sections:
                    if section.header:
                        header_text = "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
                        if header_text:
                            content.insert(0, f"HEADER: {header_text}")
                    
                    if section.footer:
                        footer_text = "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
                        if footer_text:
                            content.append(f"FOOTER: {footer_text}")
            except:
                pass  # Headers/footers might not be accessible
            
            return "\n".join(content)
            
        except ImportError:
            logger.error("python-docx not available - install with: pip install python-docx")
            return f"Error: Cannot process DOCX file - python-docx not installed"
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return f"Error processing DOCX file: {str(e)}"
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Successfully read TXT file with {encoding} encoding")
                        return content
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try binary mode
            with open(file_path, 'rb') as file:
                content = file.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return f"Error reading text file: {str(e)}"
    
    async def _extract_rtf_text(self, file_path: str) -> str:
        """Extract text from RTF file"""
        try:
            # Try to use striprtf if available
            try:
                from striprtf.striprtf import rtf_to_text
                with open(file_path, 'r', encoding='utf-8') as file:
                    rtf_content = file.read()
                return rtf_to_text(rtf_content)
            except ImportError:
                logger.warning("striprtf not available - install with: pip install striprtf")
                # Fallback to basic text extraction
                return await self._extract_txt_text(file_path)
        except Exception as e:
            logger.error(f"Error processing RTF: {e}")
            return f"Error processing RTF file: {str(e)}"
    
    async def _extract_generic_text(self, file_path: str) -> str:
        """Try to extract text from unknown file types"""
        try:
            # Try reading as text with different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            return content
                except UnicodeDecodeError:
                    continue
            
            return f"Unable to extract text from {os.path.splitext(file_path)[1]} file"
            
        except Exception as e:
            logger.error(f"Error processing generic file: {e}")
            return f"Error processing file: {str(e)}"
    
    async def _extract_document_entities(self, content: str) -> Dict:
        """Extract entities from document content"""
        if not content:
            return {
                'legal_terms': [], 'case_numbers': [], 'dates': [], 'names': [],
                'organizations': [], 'monetary_amounts': [], 'addresses': []
            }
        
        entities = {
            'legal_terms': [],
            'case_numbers': [],
            'dates': [],
            'names': [],
            'organizations': [],
            'monetary_amounts': [],
            'addresses': []
        }
        
        # Legal terms
        legal_terms = [
            'contract', 'agreement', 'lawsuit', 'plaintiff', 'defendant',
            'attorney', 'lawyer', 'counsel', 'court', 'judge', 'settlement',
            'damages', 'liability', 'negligence', 'breach', 'violation',
            'deposition', 'testimony', 'evidence', 'witness', 'objection'
        ]
        
        content_lower = content.lower()
        for term in legal_terms:
            if term in content_lower:
                entities['legal_terms'].append(term)
        
        # Extract using regex patterns
        import re
        
        # Case numbers
        case_patterns = [
            r'\b\d{2,4}-\d{2,6}\b',
            r'\bCase No\.?\s*\d+\b',
            r'\bCivil No\.?\s*\d+\b',
            r'\bDocket No\.?\s*\d+\b'
        ]
        
        for pattern in case_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            entities['case_numbers'].extend(matches)
        
        # Dates
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            entities['dates'].extend(matches)
        
        # Monetary amounts
        money_patterns = [
            r'\$[\d,]+\.?\d*',
            r'\b\d+\.\d{2}\s*dollars?\b'
        ]
        
        for pattern in money_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            entities['monetary_amounts'].extend(matches)
        
        # Names (basic pattern)
        name_pattern = r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
        potential_names = re.findall(name_pattern, content)
        
        # Filter out common non-names
        common_non_names = {
            'United States', 'New York', 'Los Angeles', 'Court Order',
            'Case Number', 'Legal Notice', 'Page Number', 'Supreme Court'
        }
        
        entities['names'] = [name for name in potential_names[:10] if name not in common_non_names]
        
        # Remove duplicates
        for key in entities:
            entities[key] = list(set(entities[key]))
        
        return entities
    
    async def _scan_document_privilege(self, content: str) -> List[Dict]:
        """Scan document for privilege indicators"""
        if not content:
            return []
        
        from app.utils.privilege_patterns import PrivilegeScanner
        scanner = PrivilegeScanner()
        return scanner.scan_text(content)
    
    async def _analyze_document_structure(self, content: str, file_format: str) -> Dict:
        """Analyze document structure and format"""
        structure = {
            'document_type': 'unknown',
            'has_header': False,
            'has_footer': False,
            'has_signature_block': False,
            'has_letterhead': False,
            'paragraph_count': 0,
            'estimated_pages': 1,
            'language': 'english'
        }
        
        if not content:
            return structure
        
        lines = content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        structure['paragraph_count'] = len(non_empty_lines)
        structure['estimated_pages'] = max(1, len(content) // 3000)  # Rough estimate
        
        content_lower = content.lower()
        
        # Document type detection
        if any(word in content_lower for word in ['contract', 'agreement', 'terms and conditions']):
            structure['document_type'] = 'contract'
        elif any(word in content_lower for word in ['memorandum', 'memo', 'subject:']):
            structure['document_type'] = 'memorandum'
        elif any(word in content_lower for word in ['dear', 'sincerely', 'yours truly']):
            structure['document_type'] = 'letter'
        elif any(word in content_lower for word in ['invoice', 'bill', 'amount due']):
            structure['document_type'] = 'invoice'
        elif any(word in content_lower for word in ['deposition', 'sworn testimony']):
            structure['document_type'] = 'deposition'
        elif any(word in content_lower for word in ['email', 'from:', 'to:', 'subject:']):
            structure['document_type'] = 'email'
        
        # Structure detection
        if len(non_empty_lines) > 0:
            first_lines = ' '.join(non_empty_lines[:3]).lower()
            if any(word in first_lines for word in ['company', 'corporation', 'llc', 'law firm']):
                structure['has_letterhead'] = True
                structure['has_header'] = True
        
        if any(word in content_lower for word in ['signature', 'signed by', 'date signed']):
            structure['has_signature_block'] = True
        
        # Page indicators
        if '--- page ' in content_lower:
            page_count = content_lower.count('--- page ')
            structure['estimated_pages'] = max(structure['estimated_pages'], page_count)
        
        return structure
    
    async def _generate_document_summary(self, content: str, entities: Dict, structure: Dict, file_format: str) -> str:
        """Generate comprehensive document summary"""
        try:
            summary_parts = []
            
            # Basic info
            word_count = len(content.split()) if content else 0
            char_count = len(content) if content else 0
            
            summary_parts.append(f"Document: {file_format.upper()}, {word_count} words")
            
            # Document type
            doc_type = structure.get('document_type', 'unknown')
            if doc_type != 'unknown':
                summary_parts.append(f"Type: {doc_type}")
            
            # Page count
            pages = structure.get('estimated_pages', 1)
            if pages > 1:
                summary_parts.append(f"Pages: {pages}")
            
            # Key entities
            if entities.get('legal_terms'):
                summary_parts.append(f"Legal terms: {len(entities['legal_terms'])}")
            
            if entities.get('case_numbers'):
                summary_parts.append(f"Case numbers: {', '.join(entities['case_numbers'][:2])}")
            
            if entities.get('monetary_amounts'):
                summary_parts.append(f"Financial references: {len(entities['monetary_amounts'])}")
            
            if entities.get('dates'):
                summary_parts.append(f"Dates mentioned: {len(entities['dates'])}")
            
            # Structure features
            if structure.get('has_signature_block'):
                summary_parts.append("Contains signatures")
            
            if structure.get('has_letterhead'):
                summary_parts.append("Has letterhead")
            
            return ". ".join(summary_parts)
            
        except Exception as e:
            logger.error(f"Error generating document summary: {e}")
            return f"Document processed successfully ({file_format.upper()})"